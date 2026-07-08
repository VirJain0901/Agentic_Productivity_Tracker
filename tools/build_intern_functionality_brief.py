from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("documentation/intern_functionality_brief.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5B677A"
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
BORDER = "C9D3DF"


def set_run(run, size=None, bold=None, color=None, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def border(cell, color=BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width * 1440)))


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def style_table(table, widths, header_fill=HEADER_FILL, font_size=8.6):
    set_table_width(table, widths)
    repeat_header(table.rows[0])
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            border(cell)
            margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_idx == 0:
                shade(cell, header_fill)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.1
                for run in paragraph.runs:
                    set_run(run, size=font_size, bold=(row_idx == 0), color=INK)


def heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    run = paragraph.add_run(text)
    if level == 1:
        set_run(run, size=16, bold=True, color=BLUE)
    elif level == 2:
        set_run(run, size=13, bold=True, color=BLUE)
    else:
        set_run(run, size=12, bold=True, color=DARK_BLUE)
    return paragraph


def para(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run(run, size=11)
    return paragraph


def bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run(run, size=10.5)
    return paragraph


def numbered(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run(run, size=10.5)
    return paragraph


def callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    shade(cell, LIGHT_FILL)
    border(cell, "D7DEE8")
    margins(cell, top=140, bottom=140, start=180, end=180)
    title_run = cell.paragraphs[0].add_run(title)
    set_run(title_run, bold=True, color=INK)
    body_p = cell.add_paragraph()
    body_p.paragraph_format.space_after = Pt(0)
    body_run = body_p.add_run(body)
    set_run(body_run, size=10.5)
    doc.add_paragraph()


def table_with_rows(doc, headers, rows, widths, font_size=8.6):
    table = doc.add_table(rows=1, cols=len(headers))
    for idx, header in enumerate(headers):
        table.cell(0, idx).text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    style_table(table, widths, font_size=font_size)
    doc.add_paragraph()


def code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    shade(cell, "F7F9FC")
    border(cell, "D8E0EA")
    margins(cell, top=120, bottom=120, start=160, end=160)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(lines):
        if idx:
            paragraph.add_run("\n")
        run = paragraph.add_run(line)
        set_run(run, size=9, color="27364A", font="Consolas")
    doc.add_paragraph()


def configure_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    return doc


def add_module(doc, module):
    heading(doc, f"{module['number']}. {module['title']} - {module['assignee']}", 1)
    callout(doc, "Objective", module["objective"])

    heading(doc, "Primary Responsibilities", 2)
    for item in module["responsibilities"]:
        bullet(doc, item)

    heading(doc, "Inputs", 2)
    for item in module["inputs"]:
        bullet(doc, item)

    heading(doc, "Outputs", 2)
    for item in module["outputs"]:
        bullet(doc, item)

    heading(doc, "Interfaces And Contracts", 2)
    table_with_rows(
        doc,
        ["Interface", "Direction", "Expected Behavior"],
        module["interfaces"],
        [1.7, 1.1, 3.7],
        font_size=8.8,
    )

    heading(doc, "Implementation Notes", 2)
    for item in module["implementation"]:
        bullet(doc, item)

    heading(doc, "Error Handling", 2)
    for item in module["errors"]:
        bullet(doc, item)

    heading(doc, "Testing Criteria", 2)
    for item in module["tests"]:
        bullet(doc, item)


MODULES = [
    {
        "number": "Module 1",
        "title": "Agent & Activity Monitoring",
        "assignee": "Bhumika",
        "objective": "Build the Python desktop capture layer that observes user/system activity and emits clean structured events for the sync client.",
        "responsibilities": [
            "Detect the active foreground application and process name at a regular interval.",
            "Capture the foreground window title for context while avoiding unnecessary sensitive content in logs.",
            "Track keyboard/mouse activity signals enough to determine active versus idle state.",
            "Detect idle periods using the configured threshold, with a default target of 600 seconds.",
            "Capture screenshot metadata at configured intervals if enabled by policy; binary streaming is not part of this assignment.",
            "Emit structured activity, idle, heartbeat, and screenshot metadata dictionaries for Module 2.",
        ],
        "inputs": [
            "Operating-system foreground window state.",
            "Running process metadata.",
            "Keyboard and mouse activity signals.",
            "Local policy configuration such as idle threshold and screenshot interval.",
        ],
        "outputs": [
            "Activity events containing timestamp, app name, process name, and window title.",
            "Idle events containing start time, end time, and total idle seconds.",
            "Screenshot metadata containing timestamp and local reference metadata when policy allows.",
            "Health or heartbeat events indicating agent loop status.",
        ],
        "interfaces": [
            ["send_activity_data(activity_dict)", "To Module 2", "Called whenever active app/window state changes or at the configured interval."],
            ["send_idle_data(idle_dict)", "To Module 2", "Called when an idle period starts/ends or when a complete idle event can be reported."],
            ["send_screenshot_metadata(meta_dict)", "To Module 2", "Called only for metadata and only when screenshot capture is enabled by policy."],
        ],
        "implementation": [
            "Keep the module independent from Django models, database code, and unrelated presentation layers.",
            "Prefer a small capture loop with clear timing controls instead of scattered background tasks.",
            "Use consistent ISO 8601 timestamps and stable event field names.",
            "Avoid counting the same idle period twice; model idle transitions explicitly.",
            "Keep CPU usage below the project target of 5 percent during normal monitoring.",
        ],
        "errors": [
            "If active window detection fails, log a structured warning and emit a safe fallback only if needed.",
            "If screenshot metadata capture fails, do not crash the agent loop; record the error and continue activity capture.",
            "If OS permissions are missing, report a clear permission error for Module 4/security handling.",
        ],
        "tests": [
            "Correctly identifies active applications while the user switches between multiple windows.",
            "Calculates idle duration accurately across active -> idle -> active transitions.",
            "Runs for an extended period without memory growth or crashes.",
            "Produces dictionaries that Module 2 can store without transformation hacks.",
        ],
    },
    {
        "number": "Module 2",
        "title": "Storage & Sync Client",
        "assignee": "Vaidehi",
        "objective": "Build the reliable Python buffering and sync layer so monitoring data is not lost when the backend or network is unavailable.",
        "responsibilities": [
            "Accept events from Module 1 and persist them locally before attempting network sync.",
            "Assign or preserve UUID event IDs and idempotency keys for every record.",
            "Maintain explicit queue states: pending, in progress, synced, retrying, and dead letter.",
            "Batch pending records for backend API submission.",
            "Parse backend sync acknowledgements as accepted, duplicate, or rejected.",
            "Retry recoverable failures using backoff and keep rejected records visible for debugging.",
        ],
        "inputs": [
            "Activity, idle, heartbeat, and screenshot metadata dictionaries from Module 1.",
            "Backend API base URL and authentication credentials.",
            "Network availability and backend response state.",
        ],
        "outputs": [
            "Durable local queue records.",
            "Batched REST API requests to Module 3.",
            "Sync status logs and dead-letter records for rejected events.",
        ],
        "interfaces": [
            ["store_record(data_dict)", "From Module 1", "Persists incoming data locally before any network attempt."],
            ["sync_pending_records()", "To Module 3", "Sends batched pending records and updates local state from backend ack response."],
            ["get_sync_status()", "For diagnostics", "Returns queue depth, retry counts, last success, and dead-letter count."],
        ],
        "implementation": [
            "Never mark a record synced merely because the HTTP request returned 207 or 200; inspect the response body.",
            "Treat duplicate acknowledgements from the backend as successful sync outcomes.",
            "Move rejected records to a dead-letter state with the backend error reason.",
            "Do not directly call Module 1 capture logic or Module 3 database internals.",
            "Keep local storage schema small, durable, and easy to inspect during testing.",
        ],
        "errors": [
            "Network failure should leave records pending or retrying, not dropped.",
            "Malformed backend responses should be logged and treated as sync failure.",
            "Database/local file errors should be surfaced clearly and not swallowed silently.",
        ],
        "tests": [
            "Events survive process restart before sync.",
            "Offline events sync automatically after network recovery.",
            "Accepted and duplicate events become synced.",
            "Rejected events move to dead letter with reason and retry metadata.",
            "No duplicate records are transmitted after retries.",
        ],
    },
    {
        "number": "Module 3",
        "title": "Backend & Real-Time System",
        "assignee": "Sanskruti",
        "objective": "Build the Python/Django backend that receives validated data, stores canonical records, and serves real-time status safely.",
        "responsibilities": [
            "Implement REST APIs for activity logs, idle events, heartbeat, screenshot metadata, policy fetch, and sync ingestion.",
            "Validate all incoming event envelopes and reject malformed records with clear errors.",
            "Store accepted records in the canonical backend database.",
            "Enforce authentication for users and devices.",
            "Publish real-time updates through tenant/team/class scoped WebSocket channels.",
            "Manage session lifecycle including session start, end, active duration, and idle duration.",
        ],
        "inputs": [
            "Batched event payloads from Module 2.",
            "Authentication tokens or device credentials.",
            "Policy requests and heartbeat updates from enrolled clients.",
        ],
        "outputs": [
            "Database records for activity, idle periods, sessions, heartbeat, and policy acknowledgement.",
            "Sync acknowledgement response containing accepted, duplicates, and rejected lists.",
            "WebSocket status updates for authorized dashboard consumers.",
            "API responses for reporting and monitoring views.",
        ],
        "interfaces": [
            ["POST /api/activity-sync/", "From Module 2", "Validates and stores event batches; returns accepted, duplicate, and rejected event IDs."],
            ["POST /api/heartbeat/", "From Module 2", "Updates device health, last-seen status, queue depth, and sync diagnostics."],
            ["GET /api/monitoring/policies/", "To clients", "Returns policy configuration scoped to the authenticated tenant/device."],
            ["WebSocket status channel", "To dashboard/API clients", "Broadcasts only authorized, tenant-scoped realtime updates."],
        ],
        "implementation": [
            "Django-managed backend storage is the source of truth.",
            "Use serializers/validators for request shape, required fields, UUIDs, and timestamps.",
            "Keep tenant/device/user boundaries explicit in every query and channel.",
            "Do not accept arbitrary client-submitted device identity without credential validation.",
            "Keep WebSocket payloads small and versioned.",
        ],
        "errors": [
            "Return 400 for invalid data with per-record reasons when possible.",
            "Return 401/403 for missing or invalid credentials.",
            "Return partial acknowledgement for mixed success/failure batches without losing rejected details.",
            "Log server errors with enough context for debugging but without exposing secrets.",
        ],
        "tests": [
            "API validation accepts correct payloads and rejects malformed payloads.",
            "Duplicate event IDs are idempotent.",
            "Unauthorized requests are denied.",
            "Tenant-scoped realtime channels do not leak data across groups.",
            "Session lifecycle calculations are correct.",
        ],
    },
    {
        "number": "Module 4",
        "title": "Security & System Persistence",
        "assignee": "Kiara",
        "objective": "Build Python/system-level reliability and policy enforcement support while keeping high-risk behavior controlled and auditable.",
        "responsibilities": [
            "Implement start-on-boot and watchdog behavior for the desktop agent.",
            "Restart the monitoring process if it exits unexpectedly.",
            "Maintain secure local configuration and authentication token handling.",
            "Fetch and apply backend policy updates.",
            "Support restricted website enforcement where policy allows.",
            "Log permission failures, policy sync failures, and service health events.",
        ],
        "inputs": [
            "Backend policies and authentication data.",
            "Operating-system service/persistence configuration.",
            "Local agent process health state.",
        ],
        "outputs": [
            "Running agent/service process.",
            "Applied local policy state.",
            "Security and watchdog logs.",
            "Permission or enforcement failure diagnostics.",
        ],
        "interfaces": [
            ["policy_sync()", "From Module 3", "Fetches policy updates and passes allowed settings to the local enforcement layer."],
            ["watchdog_loop()", "Supports Module 1", "Checks that the monitoring process is running and restarts it if necessary."],
            ["secure_config_read()", "Supports all modules", "Provides local configuration without exposing secrets in logs."],
        ],
        "implementation": [
            "Keep system persistence code separate from business logic and analytics logic.",
            "Prefer absolute paths and explicit service configuration over fragile relative command launches.",
            "Do not hard-code secrets or print tokens in logs.",
            "All production communication assumptions should be HTTPS/WSS.",
            "High-risk features such as screenshot streaming, file surveillance, risk scoring, and remote commands remain gated until legal/privacy review.",
        ],
        "errors": [
            "If the watchdog cannot restart the agent, log the failure with path, exit state, and permission hints.",
            "If policy application fails, keep the previous known-good policy where safe and report the failure.",
            "If required permissions are missing, report a structured error instead of crashing silently.",
        ],
        "tests": [
            "Agent starts automatically after reboot or simulated service start.",
            "Watchdog restarts a terminated agent process.",
            "Policy sync retries after backend/network failure.",
            "Sensitive tokens are not printed in logs.",
            "Website blocking or policy enforcement can be enabled and rolled back safely in test mode.",
        ],
    },
    {
        "number": "Module 5",
        "title": "Machine Learning & Analytics",
        "assignee": "Vikrant",
        "objective": "Build Python analytics logic that produces explainable productivity insights from approved backend data without affecting the capture/sync pipeline.",
        "responsibilities": [
            "Process historical activity/session data exported or served by Module 3.",
            "Classify applications/sites into productive, neutral, or unproductive categories using explainable rules first.",
            "Generate productivity summaries per person, session, group, and time window.",
            "Detect simple anomalies only after baseline scoring is stable.",
            "Return analytics output as safe DTOs for backend/dashboard consumption.",
        ],
        "inputs": [
            "Processed backend activity/session data.",
            "Approved app/site category mappings.",
            "Time windows and grouping criteria for reports.",
        ],
        "outputs": [
            "Productivity score and category summaries.",
            "App/site usage breakdowns.",
            "Group-level trends and report-ready aggregates.",
            "Explainability notes showing why a score/category was assigned.",
        ],
        "interfaces": [
            ["generate_productivity_summary(records)", "From Module 3 data", "Returns scored summaries without mutating backend records."],
            ["classify_activity(app_name, site)", "Internal analytics", "Maps activity to a category using approved rules or models."],
            ["export_report_data(filters)", "To backend/reporting", "Returns report-safe aggregates for display/export."],
        ],
        "implementation": [
            "Start with deterministic rule-based scoring before ML models.",
            "Keep analytics outside the live capture loop so it cannot slow monitoring.",
            "Use clear category definitions and make scoring explainable.",
            "Avoid labels such as high-risk user unless legal/privacy review explicitly approves them.",
            "Do not read private local files or perform file-system surveillance as part of this module.",
        ],
        "errors": [
            "If input data is incomplete, return partial analytics with warnings instead of crashing.",
            "If category mapping is unknown, classify as neutral/uncategorized and log for review.",
            "If model artifacts are missing, fall back to rule-based logic.",
        ],
        "tests": [
            "Known app/site examples map to expected categories.",
            "Scores are consistent for identical input data.",
            "Large datasets process without blocking backend operations.",
            "Outputs include explanation fields.",
            "Analytics never changes raw monitoring records.",
        ],
    },
]


def build_doc():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = configure_document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Intern Functionality Brief")
    set_run(run, size=22, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    run = subtitle.add_run("Employee Productivity Monitoring System - Python Module Responsibilities")
    set_run(run, size=11.5, color=MUTED)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(12)
    run = meta.add_run("Prepared for Bhumika, Vaidehi, Sanskruti, Kiara, and Vikrant. Updated 20 June 2026.")
    set_run(run, size=10, color=MUTED)

    callout(
        doc,
        "Purpose",
        "This document is only for intern module execution. Each intern should focus on the assigned Python/backend/system functionality described below.",
    )

    heading(doc, "Assignment Summary", 1)
    table_with_rows(
        doc,
        ["Module", "Assignee", "Primary Function"],
        [
            ["Agent & Activity Monitoring", "Bhumika", "Capture activity, idle state, and screenshot metadata from the system."],
            ["Storage & Sync Client", "Vaidehi", "Persist data locally and reliably transmit it to backend APIs."],
            ["Backend & Real-Time System", "Sanskruti", "Process, store, authenticate, and serve data with realtime updates."],
            ["Security & System Persistence", "Kiara", "Keep the agent running and enforce approved local policies securely."],
            ["Machine Learning & Analytics", "Vikrant", "Generate explainable productivity insights from approved backend data."],
        ],
        [2.25, 1.1, 3.15],
        font_size=8.8,
    )

    heading(doc, "Shared Rules For All Interns", 1)
    numbered(doc, "Work only inside the assigned module boundaries unless integration review approves a cross-module change.")
    numbered(doc, "Use structured JSON-like dictionaries and documented interfaces for module communication.")
    numbered(doc, "Write readable Python with small functions, clear names, and meaningful error logging.")
    numbered(doc, "Add unit tests for the module before integration testing.")
    numbered(doc, "Do not implement legal-gated features such as screenshot streaming, file surveillance, remote commands, or ML risk labels in the production path.")

    heading(doc, "Common Event Shape", 1)
    code_block(
        doc,
        [
            "{",
            "  'schema_version': '1.0',",
            "  'event_id': 'uuid',",
            "  'event_type': 'activity | idle | heartbeat | screenshot_metadata',",
            "  'occurred_at': 'ISO8601 timestamp',",
            "  'captured_at': 'ISO8601 timestamp',",
            "  'payload': { ... }",
            "}",
        ],
    )

    for module in MODULES:
        add_module(doc, module)

    heading(doc, "Integration Checklist", 1)
    table_with_rows(
        doc,
        ["Check", "Expected Result"],
        [
            ["Module 1 -> Module 2", "Activity/idle/screenshot metadata events are accepted and stored without field-name rewrites."],
            ["Module 2 -> Module 3", "Offline queue syncs through API batches and handles accepted/duplicate/rejected responses."],
            ["Module 3 -> Consumers", "Backend exposes authenticated APIs and scoped realtime events without cross-tenant leakage."],
            ["Module 4 -> Local system", "Agent persistence, watchdog, secure config, and policy enforcement work without destabilizing the machine."],
            ["Module 5 -> Backend/reporting", "Analytics runs on approved backend data and returns explainable summaries."],
        ],
        [2.1, 4.4],
        font_size=8.8,
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Intern functionality brief - Python modules only")
    set_run(run, size=8.5, color=MUTED)

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_doc().resolve())
