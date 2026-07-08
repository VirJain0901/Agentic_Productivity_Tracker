from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("documentation/intern_python_module_brief_investor_demo_track.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5B677A"
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
WARN_FILL = "FFF8E6"
BORDER = "C9D3DF"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def border(cell, color=BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run(run, size=None, bold=None, color=None, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width * 1440)))


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_table(table, widths, header_fill=HEADER_FILL, font_size=8.8):
    set_table_width(table, widths)
    repeat_header(table.rows[0])
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            border(cell)
            margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_idx == 0:
                shade(cell, header_fill)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.1
                for run in paragraph.runs:
                    set_run(run, size=font_size, bold=(row_idx == 0), color=INK)


def heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    run = paragraph.add_run(text)
    if level == 1:
        set_run(run, size=16, bold=True, color=BLUE)
    elif level == 2:
        set_run(run, size=13, bold=True, color=BLUE)
    else:
        set_run(run, size=12, bold=True, color=DARK_BLUE)
    return paragraph


def para(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run(run, size=11)
    return paragraph


def bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run(run, size=10.5)
    return paragraph


def number(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run(run, size=10.5)
    return paragraph


def callout(doc, title, body, fill=LIGHT_FILL):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    shade(cell, fill)
    border(cell, "D7DEE8")
    margins(cell, top=140, bottom=140, start=180, end=180)
    title_run = cell.paragraphs[0].add_run(title)
    set_run(title_run, bold=True, color=INK)
    body_p = cell.add_paragraph()
    body_p.paragraph_format.space_after = Pt(0)
    body_run = body_p.add_run(body)
    set_run(body_run, size=10.5)
    doc.add_paragraph()


def table_with_rows(doc, headers, rows, widths, header_fill=HEADER_FILL, font_size=8.8):
    table = doc.add_table(rows=1, cols=len(headers))
    for idx, header in enumerate(headers):
        table.cell(0, idx).text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    style_table(table, widths, header_fill=header_fill, font_size=font_size)
    doc.add_paragraph()
    return table


def code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    shade(cell, "F7F9FC")
    border(cell, "D8E0EA")
    margins(cell, top=120, bottom=120, start=160, end=160)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(lines):
        if idx:
            paragraph.add_run("\n")
        run = paragraph.add_run(line)
        set_run(run, size=9, color="27364A", font="Consolas")
    doc.add_paragraph()


def configure_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    return doc


def build_doc():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = configure_doc()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Intern Module Brief")
    set_run(run, size=22, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    run = subtitle.add_run("Python Work Assignments + Separate Investor Demo UI Track")
    set_run(run, size=11.5, color=MUTED)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(12)
    run = meta.add_run("Employee Productivity Monitoring System | Updated 20 June 2026")
    set_run(run, size=10, color=MUTED)

    callout(
        doc,
        "Ownership correction",
        "The interns are assigned to Python-based system modules. They are not responsible for converting the investor demo HTML into TypeScript. The TypeScript/Web3/GSAP/Lenis investor demo work is a separate project UI track handled by Codex / product engineering.",
        fill="EEF3FA",
    )

    heading(doc, "1. Current Direction", 1)
    bullet(doc, "Keep intern work focused on Python modules: agent capture, local storage/sync, Django backend/realtime, system security/persistence, and Python analytics.")
    bullet(doc, "Keep module boundaries strict. Interns should integrate through JSON/API/WebSocket contracts, not by reaching into each other's code.")
    bullet(doc, "The investor demo frontend will be rebuilt separately in TypeScript with Web3, GSAP, and Lenis.")
    bullet(doc, "Interns only need to provide clean Python interfaces and backend contracts that the TypeScript demo can consume later.")

    heading(doc, "2. Python Module Assignment Matrix", 1)
    table_with_rows(
        doc,
        ["#", "Module", "Assignee", "Python Responsibility", "Integration Boundary"],
        [
            ["1", "Agent & Activity Monitoring", "Bhumika", "Capture active app/window, idle time, keyboard/mouse activity signals, and screenshot metadata using Python.", "Outputs structured event dictionaries only. No direct Django ORM, TypeScript, or dashboard work."],
            ["2", "Storage & Sync Client", "Vaidehi", "Persist events locally, maintain UUID-based unsynced queue, retry failed sync, and send batches to backend APIs.", "Receives Module 1 events and sends only to Module 3 APIs. No frontend ownership."],
            ["3", "Backend & Real-Time System", "Sanskruti", "Build Python/Django APIs, validation, database storage, JWT auth, WebSocket status updates, and session tracking.", "Owns API/WebSocket contracts consumed by UI. No investor demo animation work."],
            ["4", "Security & System Persistence", "Kiara", "Implement Python/system-level persistence, watchdog behavior, policy sync, secure token handling, HTTPS/WSS assumptions, and enforcement safeguards.", "Supports modules through secure configuration and policies. No UI/Web3 implementation ownership."],
            ["5", "Machine Learning & Analytics", "Vikrant", "Build Python analytics logic for productivity categories, scoring prototypes, and future insight extraction.", "Analytics output should be exposed as backend-safe DTOs only after review. No frontend scoring logic."],
        ],
        [0.35, 1.45, 0.8, 2.25, 1.65],
        font_size=8.4,
    )

    heading(doc, "3. Separate Investor Demo UI Track", 1)
    para(doc, "This track is not assigned to the interns. It is the product/investor demo layer and should be handled separately so interns can keep their Python modules clean.")
    table_with_rows(
        doc,
        ["Area", "Owner", "Scope", "Important Boundary"],
        [
            ["HTML to TypeScript", "Codex / Product UI Track", "Convert the investor demo into a TypeScript app with typed models, reusable components, and API-ready state.", "Do not modify intern Python files to make the demo work."],
            ["Web3", "Codex / Product UI Track", "Use only for optional investor-demo identity, signed admin approval concepts, or non-sensitive audit proof concepts.", "No employee/student/activity/screenshot data on-chain."],
            ["GSAP", "Codex / Product UI Track", "Use for page transitions, modal movement, dashboard counters, live-card updates, and polished investor-demo motion.", "No animation requirements assigned to interns."],
            ["Lenis", "Codex / Product UI Track", "Use for smooth long-page scrolling in the demo while preserving keyboard navigation and reduced-motion support.", "No scroll/UX library work assigned to interns."],
        ],
        [1.3, 1.45, 2.55, 1.2],
        header_fill="F2F4F7",
        font_size=8.5,
    )

    heading(doc, "4. Per-Intern Detailed Scope", 1)

    heading(doc, "Bhumika - Agent & Activity Monitoring", 2)
    bullet(doc, "Implement Python-based foreground window detection, process name extraction, window title capture, idle state detection, and configurable screenshot metadata capture.")
    bullet(doc, "Emit structured activity, idle, and screenshot metadata events every agreed interval.")
    bullet(doc, "Keep CPU usage low and avoid blocking user workflow.")
    bullet(doc, "Testing: multitasking detection accuracy, idle-duration correctness, screenshot metadata consistency, crash-free long run.")

    heading(doc, "Vaidehi - Storage & Sync Client", 2)
    bullet(doc, "Implement Python local persistence with UUID event IDs, pending/synced/dead-letter states, and durable restart recovery.")
    bullet(doc, "Batch events to backend REST APIs and parse accepted, duplicate, and rejected acknowledgements correctly.")
    bullet(doc, "Use retry/backoff and preserve events during offline periods.")
    bullet(doc, "Testing: no duplicate sync, no data loss, recovery after offline mode, rejection stays visible.")

    heading(doc, "Sanskruti - Backend & Real-Time System", 2)
    bullet(doc, "Implement Django/Python ingestion APIs for activity, idle, screenshot metadata, heartbeat, policy fetch, and session lifecycle.")
    bullet(doc, "Validate payloads, enforce auth, store canonical records, and publish tenant-scoped realtime updates.")
    bullet(doc, "Produce API and WebSocket contracts that the TypeScript investor demo can consume later.")
    bullet(doc, "Testing: auth denial, validation, database integrity, session lifecycle, realtime tenant scoping.")

    heading(doc, "Kiara - Security & System Persistence", 2)
    bullet(doc, "Implement secure token handling, HTTPS/WSS assumptions, watchdog/service persistence, policy sync, and system enforcement safeguards.")
    bullet(doc, "Keep legal-gated features gated: no screenshot streaming, remote commands, file surveillance, or risk scoring in production without review.")
    bullet(doc, "Document what requires admin/system permissions and how failures are logged.")
    bullet(doc, "Testing: restart reliability, policy sync retry, token protection, permission failure handling.")

    heading(doc, "Vikrant - Machine Learning & Analytics", 2)
    bullet(doc, "Implement Python analytics prototypes using backend-approved historical data only.")
    bullet(doc, "Start with explainable rule-based categorization before any advanced ML.")
    bullet(doc, "Expose analytics as safe backend DTOs; do not put scoring logic in the frontend.")
    bullet(doc, "Testing: consistent scoring, known app categorization, explainability, no backend performance impact.")

    heading(doc, "5. Shared Python Contracts", 1)
    code_block(
        doc,
        [
            "activity_event = {",
            "  'timestamp': 'ISO8601',",
            "  'app_name': 'chrome.exe',",
            "  'window_title': 'Research notes',",
            "  'event_id': 'uuid',",
            "}",
            "",
            "sync_ack = {",
            "  'accepted': ['uuid'],",
            "  'duplicates': ['uuid'],",
            "  'rejected': [{'event_id': 'uuid', 'error': 'reason'}],",
            "}",
        ],
    )
    bullet(doc, "Module 1 sends events only to Module 2.")
    bullet(doc, "Module 2 is the integration bridge to Module 3.")
    bullet(doc, "Module 3 is the backend source of truth.")
    bullet(doc, "UI work consumes Module 3 contracts later; it does not become intern responsibility.")

    heading(doc, "6. What Interns Should Not Do", 1)
    table_with_rows(
        doc,
        ["Do Not Assign To Intern Python Modules", "Reason"],
        [
            ["TypeScript investor demo conversion", "Handled separately by Codex / Product UI Track."],
            ["GSAP animation implementation", "Frontend polish concern, not Python module work."],
            ["Lenis scroll implementation", "Frontend shell concern, not Python system module work."],
            ["Web3 wallet/signature UI", "Investor-demo feature and security review item, not intern Python scope."],
            ["Screenshot streaming or remote command execution", "Legal/privacy/audit gated feature, not safe for intern production path."],
        ],
        [3.2, 3.3],
        header_fill=WARN_FILL,
        font_size=8.7,
    )

    heading(doc, "7. Development Rules", 1)
    number(doc, "Each intern develops independently inside their assigned Python module.")
    number(doc, "All communication between modules must use documented interfaces, JSON payloads, REST APIs, or WebSocket contracts.")
    number(doc, "No intern should change another intern's module without review.")
    number(doc, "Unit tests are required before integration.")
    number(doc, "Integration tests are required before demo wiring.")
    number(doc, "The TypeScript investor demo must adapt to the Python/backend contracts, not the other way around.")

    heading(doc, "8. Definition Of Done", 1)
    bullet(doc, "Each Python module runs independently and passes its module tests.")
    bullet(doc, "The storage/sync path can survive restart, offline mode, and partial backend rejection.")
    bullet(doc, "The backend exposes documented API/WebSocket contracts.")
    bullet(doc, "Security and policy behavior is logged, testable, and gated where required.")
    bullet(doc, "Analytics outputs are explainable and safe to display.")
    bullet(doc, "The investor demo TypeScript work is tracked separately from intern Python deliverables.")

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Intern Python module brief - investor demo UI handled separately")
    set_run(run, size=8.5, color=MUTED)

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_doc().resolve())
