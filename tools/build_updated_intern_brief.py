from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("documentation/updated_intern_module_brief_typescript_web3.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
BORDER = "C9D3DF"
MUTED = "5B677A"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width * 1440)))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run(run, size=None, bold=None, color=None, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run(run)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run(run, size=10.5)
    return paragraph


def add_number(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run(run, size=10.5)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    run = paragraph.add_run(text)
    if level == 1:
        set_run(run, size=16, bold=True, color=BLUE)
    elif level == 2:
        set_run(run, size=13, bold=True, color=BLUE)
    else:
        set_run(run, size=12, bold=True, color=DARK_BLUE)
    return paragraph


def add_callout(doc, title, body, fill=CALLOUT):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color="D7DEE8")
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    paragraph = cell.paragraphs[0]
    title_run = paragraph.add_run(title)
    set_run(title_run, bold=True, color=INK)
    body_paragraph = cell.add_paragraph()
    body_paragraph.paragraph_format.space_after = Pt(0)
    body_run = body_paragraph.add_run(body)
    set_run(body_run, size=10.5)
    doc.add_paragraph()
    return table


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F9FC")
    set_cell_border(cell, color="D8E0EA")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(lines):
        if idx:
            paragraph.add_run("\n")
        run = paragraph.add_run(line)
        set_run(run, size=9, font="Consolas", color="27364A")
    doc.add_paragraph()


def style_table(table, header_fill=LIGHT_BLUE, font_size=8.7):
    set_repeat_table_header(table.rows[0])
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            set_cell_margins(cell)
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.1
                for run in paragraph.runs:
                    set_run(run, size=font_size, bold=(row_idx == 0), color=INK)


def build_doc():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run("Updated Intern Module Brief")
    set_run(run, size=22, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    run = subtitle.add_run("Employee Productivity Monitoring System | TypeScript, Web3, GSAP, and Lenis Addendum")
    set_run(run, size=11.5, color=MUTED)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(12)
    run = meta.add_run("Prepared for intern execution and project-level frontend modernization. Date: 20 June 2026.")
    set_run(run, size=10, color=MUTED)

    add_callout(
        doc,
        "Senior direction",
        "The static HTML product shell should move to a TypeScript frontend. The intern modules remain independently owned and should communicate only through agreed APIs, schemas, and WebSocket contracts. Do not rewrite intern module code just to support UI work.",
        fill="EEF3FA",
    )

    add_heading(doc, "1. Change Summary", 1)
    add_bullet(doc, "Shift the current HTML dashboard/demo surface into a TypeScript application with typed API clients and reusable UI modules.")
    add_bullet(doc, "Use GSAP for deliberate, high-quality transitions and dashboard micro-interactions.")
    add_bullet(doc, "Use Lenis for smooth page scrolling where it improves long dashboard/report pages, while preserving keyboard and reduced-motion accessibility.")
    add_bullet(doc, "Use Web3 only for approved identity/demo/audit-adjacent flows. Do not store employee/student monitoring data, screenshots, or private activity data on-chain.")
    add_bullet(doc, "Keep business logic in backend/core services. The frontend displays and controls workflows through documented APIs only.")

    add_heading(doc, "2. Updated Assignment Matrix", 1)
    table = doc.add_table(rows=1, cols=5)
    headers = ["#", "Module", "Assignee", "Primary Function", "TypeScript / UI Integration Impact"]
    for idx, header in enumerate(headers):
        table.cell(0, idx).text = header
    rows = [
        ["1", "Agent & Activity Monitoring", "Bhumika", "Capture user activity, idle time, and screenshot metadata from the system.", "Must produce structured event envelopes that can be represented as shared TypeScript types. No direct dashboard or backend ORM access."],
        ["2", "Storage & Sync Client", "Vaidehi", "Persist data locally and reliably transmit it to backend.", "Must expose sync status, queue health, UUID event IDs, accepted/duplicate/rejected acknowledgements, and dead-letter reasons for dashboard visibility."],
        ["3", "Backend & Real-Time System", "Sanskruti", "Process, store, and serve data with real-time tracking.", "Owns the APIs, OpenAPI/JSON schemas, WebSocket payloads, tenant-scoped rooms, and typed response contracts consumed by the TypeScript frontend."],
        ["4", "Security & System Persistence", "Kiara", "Ensure system runs continuously and enforce policies.", "Owns secure token handling, HTTPS/WSS assumptions, legal gates, Web3 security boundaries, CSP guidance, and system policy enforcement constraints."],
        ["5", "Machine Learning & Analytics", "Vikrant", "Generate insights and productivity intelligence.", "Must expose analytics as safe, typed DTOs through backend endpoints. No risk scoring or sensitive inference in the frontend until legal review."],
        ["6", "Frontend Dashboard & Interaction Layer", "TBD / Product UI Owner", "Migrate HTML to a production TypeScript UI with modern interactions.", "Owns TypeScript app structure, API client, WebSocket client, Web3 adapter, GSAP animation layer, Lenis scroll setup, UI states, exports, filters, and sample-data labels."],
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    set_table_width(table, [0.35, 1.35, 0.8, 1.7, 2.3])
    style_table(table)

    doc.add_paragraph()
    add_heading(doc, "3. Frontend Modernization Scope", 1)
    add_para(doc, "The existing HTML should be treated as a prototype shell. The target is a typed frontend that can be wired to real backend APIs without mixing monitoring logic into the UI.")
    add_heading(doc, "Recommended application structure", 2)
    add_code_block(
        doc,
        [
            "frontend/",
            "  src/main.ts",
            "  src/api/              # typed REST clients",
            "  src/realtime/         # WebSocket client and room subscriptions",
            "  src/types/            # generated or shared DTO types",
            "  src/web3/             # wallet/session adapter only if approved",
            "  src/animations/       # GSAP timelines and motion utilities",
            "  src/scroll/           # Lenis setup and cleanup",
            "  src/pages/            # dashboard, detail, reports, settings",
            "  src/styles/           # light/dark tokens and responsive layout",
        ],
    )
    add_bullet(doc, "Preferred baseline: Vite + TypeScript. If a framework is approved, use React + TypeScript with typed hooks and components.")
    add_bullet(doc, "The frontend must not write directly to databases, local queue files, screenshots, host files, or ML artifacts.")
    add_bullet(doc, "Sample/demo data must be labeled visibly as Sample Data until connected to authenticated APIs.")

    add_heading(doc, "4. TypeScript Contract Requirements", 1)
    add_para(doc, "The frontend should consume contracts generated from backend JSON Schema/OpenAPI wherever possible. Avoid manually inventing field names inside components.")
    add_code_block(
        doc,
        [
            "export interface ClientEventEnvelope {",
            "  schema_version: \"1.0\";",
            "  tenant_id: string;",
            "  device_id: string;",
            "  event_id: string;      // UUID",
            "  idempotency_key: string;",
            "  event_type: \"activity\" | \"idle\" | \"heartbeat\" | \"policy\";",
            "  occurred_at: string;   // ISO date-time",
            "  captured_at: string;   // ISO date-time",
            "  payload: Record<string, unknown>;",
            "}",
            "",
            "export interface SyncAck {",
            "  accepted: string[];",
            "  duplicates: string[];",
            "  rejected: Array<{ event_id: string; error: string }>; ",
            "}",
        ],
    )
    add_bullet(doc, "Every API response used by the dashboard should have an exported TypeScript type.")
    add_bullet(doc, "WebSocket messages should include tenant scope, message type, timestamp, and typed payload.")
    add_bullet(doc, "Frontend filters, reports, and exports should operate on typed API response models, not raw ad-hoc objects.")

    add_heading(doc, "5. Web3 Usage Boundary", 1)
    add_callout(doc, "Web3 rule", "Web3 must not become the storage layer for monitoring data. Use it only where it creates clear product value and has been approved by security/legal review.", fill="FFF8E6")
    add_heading(doc, "Allowed Web3 uses", 2)
    add_bullet(doc, "Wallet connection for an optional admin/investor demo identity flow.")
    add_bullet(doc, "Signature-based proof that an authorized admin approved a sensitive configuration change.")
    add_bullet(doc, "Optional audit hash pointer for non-sensitive verification metadata, after legal review.")
    add_heading(doc, "Not allowed", 2)
    add_bullet(doc, "Do not put student, employee, screenshot, browsing, activity, or productivity data on-chain.")
    add_bullet(doc, "Do not replace JWT/device credentials or backend object permissions with wallet-only checks.")
    add_bullet(doc, "Do not require monitored students/employees to use wallets.")
    add_bullet(doc, "Do not let Web3 calls block core dashboard loading or monitoring workflows.")

    add_heading(doc, "6. GSAP Animation Standards", 1)
    add_para(doc, "GSAP should make the interface feel precise and premium, not decorative or distracting.")
    add_bullet(doc, "Use GSAP for page transitions, modal/drawer entrances, timeline reveals, status changes, counters, and live-card updates.")
    add_bullet(doc, "Keep animations short: most UI transitions should land within 150-350 ms.")
    add_bullet(doc, "Respect prefers-reduced-motion and provide instant-state alternatives.")
    add_bullet(doc, "Clean up timelines when components unmount. Avoid global timelines that keep running in hidden views.")
    add_bullet(doc, "Do not animate large tables row-by-row if it harms dashboard performance.")
    add_code_block(
        doc,
        [
            "import gsap from \"gsap\";",
            "",
            "export function revealPanel(target: Element) {",
            "  return gsap.fromTo(target,",
            "    { autoAlpha: 0, y: 12 },",
            "    { autoAlpha: 1, y: 0, duration: 0.24, ease: \"power2.out\" }",
            "  );",
            "}",
        ],
    )

    add_heading(doc, "7. Lenis Scroll Standards", 1)
    add_para(doc, "Lenis should be initialized once at the app shell level and used only where smooth scrolling improves the experience.")
    add_bullet(doc, "Use Lenis on long dashboard, report, audit, or landing/demo pages.")
    add_bullet(doc, "Preserve native keyboard navigation, focus behavior, anchor links, and modal scroll locks.")
    add_bullet(doc, "Disable or simplify smooth scroll when prefers-reduced-motion is active.")
    add_bullet(doc, "Avoid double scroll containers unless the UX absolutely requires them.")
    add_code_block(
        doc,
        [
            "import Lenis from \"lenis\";",
            "",
            "export function createLenis() {",
            "  const lenis = new Lenis({ duration: 1.0, smoothWheel: true });",
            "  function raf(time: number) {",
            "    lenis.raf(time);",
            "    requestAnimationFrame(raf);",
            "  }",
            "  requestAnimationFrame(raf);",
            "  return lenis;",
            "}",
        ],
    )

    add_heading(doc, "8. UI Design Tokens To Preserve", 1)
    add_para(doc, "Use the approved neutral blue system and fonts. Edition-specific language should change labels, not the core component system.")
    colors = doc.add_table(rows=1, cols=5)
    for idx, header in enumerate(["Mode", "Background", "Surface", "Text", "Primary / Accent"]):
        colors.cell(0, idx).text = header
    for row in [
        ["Light", "#F0F4FA", "#FFFFFF", "#0C1C4E", "#1A3D8F / #C2D0EC"],
        ["Dark", "#03081E", "#071035", "#E8EEF8", "#2C5BA8 / #8AAAD8"],
    ]:
        cells = colors.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    set_table_width(colors, [0.85, 1.4, 1.25, 1.25, 1.75])
    style_table(colors, header_fill=LIGHT_GRAY, font_size=9.2)
    add_bullet(doc, "Fonts: Manrope for headings, Figtree for dashboard labels, Inter for dense tables and forms.")
    add_bullet(doc, "Cards and controls should remain compact, readable, and manager/teacher friendly rather than marketing-heavy.")
    add_bullet(doc, "Dashboard views must support both education and workforce language without forking the application.")

    add_heading(doc, "9. Intern-Specific Acceptance Checks", 1)
    checks = doc.add_table(rows=1, cols=3)
    for idx, header in enumerate(["Assignee", "Must Demonstrate", "Must Avoid"]):
        checks.cell(0, idx).text = header
    for row in [
        ["Bhumika", "Activity, idle, and screenshot metadata match the shared event envelope and are stable under multitasking.", "Direct Django ORM access, direct frontend calls, binary screenshot streaming without legal gate."],
        ["Vaidehi", "Local queue survives restart, uses UUID event IDs, parses accepted/duplicate/rejected sync responses correctly.", "Marking HTTP 207 as total success, dropping rejected records, duplicate transmissions."],
        ["Sanskruti", "APIs and WebSockets are tenant-scoped, authenticated, documented, and consumable by TypeScript clients.", "Global realtime rooms, untyped payload drift, unauthenticated dashboard data."],
        ["Kiara", "Tokens, policies, service persistence, legal gates, HTTPS/WSS, and Web3 boundaries are documented and testable.", "Storing secrets in frontend, bypassing backend auth, forcing Web3 into monitored-person flows."],
        ["Vikrant", "Analytics DTOs are versioned, explainable, and safe to display in dashboard summaries.", "Frontend-side scoring logic, high-risk labels, sensitive inference without review."],
    ]:
        cells = checks.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    set_table_width(checks, [0.9, 2.85, 2.75])
    style_table(checks, header_fill=LIGHT_BLUE, font_size=8.8)

    add_heading(doc, "10. Development Rules After This Update", 1)
    add_number(doc, "Frontend work starts in a separate TypeScript app path and consumes APIs/contracts. It must not patch intern modules directly.")
    add_number(doc, "All new UI data models must map to backend JSON Schema/OpenAPI definitions.")
    add_number(doc, "Animation and smooth scrolling must be accessible, measurable, and easy to disable.")
    add_number(doc, "Web3 must be feature-flagged and legally/security reviewed before production use.")
    add_number(doc, "No screenshot streaming, file-system surveillance, ML risk scoring, or remote commands should enter the production path until consent, audit, retention, and legal review are complete.")

    add_heading(doc, "11. Definition Of Done", 1)
    add_bullet(doc, "The HTML prototype is replaced by a TypeScript app with typed API and WebSocket clients.")
    add_bullet(doc, "The UI has empty states, loading states, error states, role-gated views, filters, reports, and export-ready screens.")
    add_bullet(doc, "GSAP animations are scoped, performant, and reduced-motion aware.")
    add_bullet(doc, "Lenis scrolling works without breaking keyboard navigation, modals, or anchor links.")
    add_bullet(doc, "Web3 is isolated behind an adapter and feature flag with no sensitive monitoring data on-chain.")
    add_bullet(doc, "Each intern module passes its own tests and integrates only through the defined contracts.")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Updated intern brief - TypeScript/Web3/GSAP/Lenis addendum")
    set_run(run, size=8.5, color=MUTED)

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_doc().resolve())
